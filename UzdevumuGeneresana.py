import random
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import matplotlib.pyplot as plt
import os

MAX_NUMBER = 50
DEFAULT_TASKS = 40


def gen_linear_equation():
    a = random.randint(1, 10)
    x = random.randint(-10, 10)
    b = a * x + random.randint(-15, 15)
    q = f"Atrisini lineāro vienādojumu: {a}x + {b - a * x} = {b}"
    return q, x


def gen_quadratic_equation():
    x1 = random.randint(-10, 10)
    x2 = random.randint(-10, 10)
    a = random.choice([1, 2, 3])
    b = -a * (x1 + x2)
    c = a * x1 * x2
    q = f"Atrisini kvadrātvienādojumu: {a}x² + {b}x + {c} = 0"
    return q, (x1, x2)


def gen_system():
    x = random.randint(-10, 10)
    y = random.randint(-10, 10)

    a1, b1 = random.randint(1, 8), random.randint(1, 8)
    a2, b2 = random.randint(1, 8), random.randint(1, 8)

    c1 = a1 * x + b1 * y
    c2 = a2 * x + b2 * y

    q = f"""Atrisini sistēmu:
{a1}x + {b1}y = {c1}
{a2}x + {b2}y = {c2}"""

    return q, (x, y)


def gen_progression():
    a1 = random.randint(-20, 20)
    d = random.randint(-10, 10)
    n = random.randint(5, 15)
    q = f"Aritmētiskā progresija: a1 = {a1}, d = {d}. Atrodi a_{n}."
    ans = a1 + (n - 1) * d
    return q, ans


def gen_geometry_triangle():
    a = random.randint(3, 15)
    b = random.randint(3, 15)
    c = random.randint(3, 15)
    q = f"Dota trijstūra malas: a={a}, b={b}, c={c}. Pārbaudi, vai trijstūris eksistē un aprēķini perimetru."
    exists = a + b > c and a + c > b and b + c > a
    ans = f"Eksistē: {exists}, Perimetrs = {a + b + c if exists else 'nav definēts'}"
    return q, ans


def gen_geometry_circle():
    r = random.randint(2, 12)
    q = f"Dots riņķis ar rādiusu r={r}. Atrodi riņķa laukumu un garumu."
    area = 3.14159 * r * r
    length = 2 * 3.14159 * r
    return q, (round(area, 3), round(length, 3))


def generate_plot_linear(path):
    x = list(range(-10, 11))
    k = random.randint(-5, 5)
    b = random.randint(-10, 10)
    y = [k * i + b for i in x]

    plt.figure(figsize=(4, 3))
    plt.plot(x, y)
    plt.title(f"y = {k}x + {b}")
    plt.grid(True)
    plt.savefig(path)
    plt.close()
    return f"Uzzīmēta funkcija y = {k}x + {b}."


def create_docx(filepath, tasks):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    doc.add_heading('Uzdevumi (9. klasei)', level=1)

    for i, (q, ans, img) in enumerate(tasks, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {q}")
        if img:
            doc.add_picture(img, width=Inches(4))
            os.remove(img)

    doc.add_page_break()
    doc.add_heading('Atbildes', level=1)
    for i, (_, ans, _) in enumerate(tasks, start=1):
        doc.add_paragraph(f"{i}. {ans}")

    doc.save(filepath)


def generate_all():
    try:
        count = int(entry_count.get())
    except:
        messagebox.showerror("Kļūda", "Ievadi derīgu skaitli!")
        return

    tasks = []

    generators = [
        gen_linear_equation,
        gen_quadratic_equation,
        gen_system,
        gen_progression,
        gen_geometry_triangle,
        gen_geometry_circle,
    ]

    for _ in range(count):
        g = random.choice(generators)
        q, ans = g()

        img = None
        if random.random() < 0.3:  
            img = f"plot_{random.randint(1,999999)}.png"
            generate_plot_linear(img)

        tasks.append((q, ans, img))

    filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX", "*.docx")])
    if not filepath:
        return

    create_docx(filepath, tasks)
    messagebox.showinfo("OK", "Fails saglabāts!")


root = tk.Tk()
root.title("Uzdevumu ģenerators 9. klasei")
root.geometry("400x200")

label = tk.Label(root, text="Cik uzdevumus ģenerēt?")
label.pack(pady=10)

entry_count = tk.Entry(root)
entry_count.insert(0, str(DEFAULT_TASKS))
entry_count.pack(pady=5)

btn = tk.Button(root, text="Ģenerēt DOCX", command=generate_all)
btn.pack(pady=20)

root.mainloop()
